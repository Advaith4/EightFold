"""DOCX loader tests for Sprint 4.3B."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pydantic import ValidationError
from src.loaders import (
    BaseLoader,
    CorruptedFileError,
    DOCXLoader,
    ExtractionStatus,
    FilePayload,
    FileReadError,
    UnsupportedFileTypeError,
)


def build_docx(*, include_text: bool = True, include_table: bool = True) -> bytes:
    """Create a minimal DOCX document for loader contract tests."""
    document = Document()
    if include_text:
        document.add_heading("Profile", level=1)
        document.add_paragraph("First paragraph")
    if include_table:
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Left cell"
        table.cell(0, 1).text = "Right cell"
    if include_text:
        document.add_paragraph("Final paragraph")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_loader_loads_valid_docx(tmp_path: Path) -> None:
    """DOCXLoader extracts deterministic document text into a FilePayload."""
    path = tmp_path / "document.docx"
    path.write_bytes(build_docx())

    payload = DOCXLoader().load(path)

    assert isinstance(payload, FilePayload)
    assert payload.text == (
        "Profile\nFirst paragraph\nLeft cell\nRight cell\nFinal paragraph"
    )
    assert payload.content_bytes is None
    assert payload.metadata.filename == "document.docx"
    assert payload.metadata.extension == ".docx"
    assert (
        payload.metadata.content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert payload.metadata.checksum.startswith("sha256:")
    assert payload.metadata.extraction_status is ExtractionStatus.TEXT_EXTRACTED


def test_docx_loader_rejects_corrupted_docx(tmp_path: Path) -> None:
    """DOCXLoader rejects malformed DOCX bytes without returning partial results."""
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a docx archive")

    with pytest.raises(CorruptedFileError):
        DOCXLoader().load(path)


def test_docx_loader_rejects_empty_docx_document(tmp_path: Path) -> None:
    """DOCXLoader rejects valid DOCX files that contain no deterministic text."""
    path = tmp_path / "empty.docx"
    path.write_bytes(build_docx(include_text=False, include_table=False))

    with pytest.raises(CorruptedFileError):
        DOCXLoader().load(path)


def test_docx_loader_rejects_missing_file(tmp_path: Path) -> None:
    """DOCXLoader raises a loader read error for missing files."""
    with pytest.raises(FileReadError):
        DOCXLoader().load(tmp_path / "missing.docx")


def test_docx_loader_rejects_unsupported_extension(tmp_path: Path) -> None:
    """DOCXLoader validates the technical file extension."""
    path = tmp_path / "document.txt"
    path.write_bytes(build_docx())

    with pytest.raises(UnsupportedFileTypeError):
        DOCXLoader().load(path)


def test_docx_loader_rejects_unreadable_file(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """DOCXLoader reports read failures through the loader exception hierarchy."""
    path = tmp_path / "document.docx"
    path.write_bytes(build_docx())

    def fail_read_bytes(self: Path) -> bytes:
        raise OSError("unreadable")

    monkeypatch.setattr(Path, "read_bytes", fail_read_bytes)

    with pytest.raises(FileReadError):
        DOCXLoader().load(path)


def test_docx_loader_rejects_empty_file(tmp_path: Path) -> None:
    """DOCXLoader rejects zero-byte files before document parsing."""
    path = tmp_path / "empty.docx"
    path.write_bytes(b"")

    with pytest.raises(CorruptedFileError):
        DOCXLoader().load(path)


def test_docx_loader_extracts_paragraphs_without_table(tmp_path: Path) -> None:
    """DOCXLoader extracts paragraph and heading text without interpretation."""
    path = tmp_path / "paragraphs.docx"
    path.write_bytes(build_docx(include_table=False))

    payload = DOCXLoader().load(path)

    assert payload.text == "Profile\nFirst paragraph\nFinal paragraph"


def test_docx_loader_extracts_table_text(tmp_path: Path) -> None:
    """DOCXLoader includes table cell text in document order."""
    path = tmp_path / "table.docx"
    path.write_bytes(build_docx(include_text=False, include_table=True))

    payload = DOCXLoader().load(path)

    assert payload.text == "Left cell\nRight cell"


def test_docx_loader_preserves_empty_paragraph_positions(tmp_path: Path) -> None:
    """DOCXLoader keeps empty paragraphs as blank lines like empty PDF text segments."""
    document = Document()
    document.add_paragraph("Before")
    document.add_paragraph("")
    document.add_paragraph("After")
    buffer = BytesIO()
    document.save(buffer)
    path = tmp_path / "empty-paragraph.docx"
    path.write_bytes(buffer.getvalue())

    payload = DOCXLoader().load(path)

    assert payload.text == "Before\n\nAfter"


def test_docx_loader_payload_serializes(tmp_path: Path) -> None:
    """DOCXLoader payloads serialize and deserialize as FilePayload objects."""
    path = tmp_path / "document.docx"
    path.write_bytes(build_docx())
    payload = DOCXLoader().load(path)

    dumped = payload.model_dump(mode="json")
    reloaded = FilePayload.model_validate(dumped)

    assert dumped["metadata"]["extraction_status"] == "text_extracted"
    assert dumped["metadata"]["extension"] == ".docx"
    assert reloaded == payload


def test_docx_loader_payload_is_immutable(tmp_path: Path) -> None:
    """DOCXLoader returns immutable loader infrastructure models."""
    path = tmp_path / "document.docx"
    path.write_bytes(build_docx())
    payload = DOCXLoader().load(path)

    with pytest.raises(ValidationError):
        payload.text = "changed"


def test_docx_loader_implements_base_loader_contract() -> None:
    """DOCXLoader is a concrete BaseLoader implementation."""
    loader = DOCXLoader()

    assert isinstance(loader, BaseLoader)
